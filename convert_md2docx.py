import sys
import re
from docx import Document
from docx.shared import Pt, Inches

def markdown_to_docx(md_path, docx_path):
    document = Document()
    
    # Configure styles
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    try:
        with open(md_path, 'r', encoding='utf-8') as f:
            content = f.read()
    except FileNotFoundError:
        print(f"Error: Could not find the file at {md_path}")
        return
        
    lines = content.split('\n')
    
    in_code_block = False
    
    for line in lines:
        if line.startswith('```'):
            in_code_block = not in_code_block
            continue
            
        if in_code_block:
            p = document.add_paragraph(line)
            # basic fixed width formatting
            for run in p.runs:
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
            continue
            
        if line.startswith('# '):
            document.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            document.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            document.add_heading(line[4:].strip(), level=3)
        elif line.startswith('#### '):
            document.add_heading(line[5:].strip(), level=4)
        elif line.startswith('> '):
            p = document.add_paragraph(line[2:].strip())
            # rudimentary quote style
        elif line.startswith('- ') or line.startswith('* '):
            p = document.add_paragraph(line[2:].strip(), style='List Bullet')
        elif line.startswith('---'):
            document.add_page_break()
        elif line.strip() == '':
            continue
        elif '|' in line and not line.startswith(' '):
            # Basic table handling (just formatting as text for simplicity)
            p = document.add_paragraph(line)
            for run in p.runs:
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
        else:
            # Handle basic bold/italic inline
            p = document.add_paragraph()
            
            # Very basic markdown inline parser
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', line)
            
            for part in parts:
                if not part:
                    continue
                    
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Consolas'
                else:
                    # Strip links logic [text](url) -> text
                    text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', part)
                    if text:
                        p.add_run(text)
    
    document.save(docx_path)
    print(f"Successfully created {docx_path}")

if __name__ == "__main__":
    # Hardcoded paths so you don't have to specify them!
    input_file = r"C:\Users\Rohith\.gemini\antigravity\brain\dfe656da-1746-4f12-afd5-9a60e5c32c04\AUTONOMOUSQA_DOCUMENTATION.md"
    output_file = r"c:\testproject\AUTONOMOUSQA_DOCUMENTATION.docx"
    
    # Still allow arguments if someone wants to override it in the future
    if len(sys.argv) == 3:
        input_file = sys.argv[1]
        output_file = sys.argv[2]
        
    print(f"Reading from: {input_file}")
    markdown_to_docx(input_file, output_file)
